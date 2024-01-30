from flask import Flask, send_file, request,render_template
from docx import Document
import os
import json

app = Flask(__name__)

@app.route('/')
def index():
    # Render the HTML form
    return render_template('main.html')

@app.route('/generate_word', methods=['POST'])
def generate_word():
    # Retrieve JSON data from the form
    json_input = request.form['jsonInput']

    try:
        # Parse the JSON data
        json_data = json.loads(json_input)
    except json.JSONDecodeError as e:
        return f"Error parsing JSON: {str(e)}"

    # Create a new Word document
    doc = Document()

    # Add content to the Word document
    doc.add_heading('JSON Data', level=1)
    for key, value in json_data.items():
        doc.add_paragraph(f"{key}: {value}")

    # Save the Word document in the current working directory
    save_path = os.path.join(os.getcwd(), 'output.docx')
    doc.save(save_path)

    # Send the Word document as a response
    return send_file(save_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
